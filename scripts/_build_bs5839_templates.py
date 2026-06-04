#!/usr/bin/env python3
"""Build the four BHO-branded BS 5839-1 certificate templates from
scratch using python-docx, mirroring the Churches Fire Security
source forms' visual style as closely as the Word OOXML model allows.

Templates produced:
  bs5839-installation-cert-template.docx   — BS 5839-1 §43.2 (A056)
  bs5839-commissioning-cert-template.docx  — BS 5839-1 §43.4 (A051),
                                              3 pages including the
                                              full 33-item §39
                                              commissioning checklist
                                              on page 2
  bs5839-acceptance-cert-template.docx     — BS 5839-1 §43.5 (A038)
  bs5839-battery-calc-template.docx        — BS 5839-1 §25.4 (A058)

Visual model (matches Churches A056-G / A051-G / A038-H / A058-G):
  - Centred logo at top of every page
  - Centred bold title in black with form code + name
  - Form content lives inside grey-shaded SECTION PANELS
  - Each panel opens with a RED section header
  - Field labels are RED on the grey panel
  - Field inputs are WHITE cells with a thin grey border (nested
    inside the panel cell)
  - Blue SUBMIT button right-aligned at the bottom of each cert
  - Sans-serif body throughout (Calibri — universal Word font)

The substitution engine in the existing edge functions fills
[Square Bracketed Placeholders] by exact <w:t> match — every
variable field in the template is a placeholder so the same
template renders correctly for any tenant. Company name +
address + logo come from company_settings; engineer + site +
customer come from the visit / report bundle.

Run with: ./scripts/build-bs5839-templates.sh
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "assets"
LOGO = ROOT / "public" / "bho-fire-logo.png"

# ── Churches palette (eyeballed off the source PDFs) ──────────────────
SECTION_GREY   = "EFF1F4"                       # the soft blue-grey panel fill
LABEL_RED      = RGBColor(0xE0, 0x3E, 0x3E)     # red section headers + field labels
INPUT_BG       = "FFFFFF"                       # white inside input cells
INPUT_BORDER   = "C7CCD3"                       # thin grey border around inputs
SUBMIT_BLUE    = "3D5AFE"                       # the bright blue of the SUBMIT button
BODY_DARK      = RGBColor(0x1A, 0x1F, 0x2C)     # near-black body text
CHECKLIST_HDR  = "1F2937"                       # dark grey for the 33-item table header
BODY_FONT      = "Calibri"                      # universal Word font

CHECKBOX_EMPTY = "☐"


# ── Low-level OOXML helpers ──────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = INPUT_BORDER, sz: int = 4) -> None:
    """Apply a thin border on all four edges of a cell — used for
    the white input cells inside the grey section panels."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_margin(cell, top: int = 80, bottom: int = 80, left: int = 100, right: int = 100) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        tcMar.append(e)
    tcPr.append(tcMar)


def _strip_table_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tblPr.append(borders)


# ── Page chrome ──────────────────────────────────────────────────────

def add_logo(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(3.8))
    else:
        run = p.add_run("[BHO LOGO]")
        run.bold = True
        run.font.color.rgb = LABEL_RED
        run.font.size = Pt(13)


def add_title(doc: Document, code: str, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"{code}  {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BODY_DARK
    run.font.name = BODY_FONT


def add_submit_footer(doc: Document) -> None:
    """Right-aligned bright-blue SUBMIT button approximation."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(3.5)
    _set_cell_shading(cell, SUBMIT_BLUE)
    _set_cell_margin(cell, top=120, bottom=120, left=240, right=240)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SUBMIT")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = BODY_FONT
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _strip_table_borders(tbl)


def add_spacer(doc: Document, pt: int = 4) -> None:
    p = doc.add_paragraph()
    p.add_run("").font.size = Pt(pt)
    p.paragraph_format.space_after = Pt(0)


# ── Section panels ───────────────────────────────────────────────────

def section_open(doc: Document, title: str | None = None) -> "object":
    """Open a grey-shaded section panel. Returns the inner cell so
    callers add field rows directly into it. Pass `title=None` to
    skip the header row (e.g. continuation panels)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(17.5)
    _set_cell_shading(cell, SECTION_GREY)
    _set_cell_margin(cell, top=160, bottom=160, left=200, right=200)
    _strip_table_borders(tbl)
    # First paragraph in the cell — replaced by the section title
    # when provided.
    first = cell.paragraphs[0]
    if title is not None:
        run = first.add_run(title)
        run.bold = True
        run.font.color.rgb = LABEL_RED
        run.font.size = Pt(10)
        run.font.name = BODY_FONT
        first.paragraph_format.space_after = Pt(6)
    else:
        first.add_run("").font.size = Pt(2)
    return cell


def section_body_paragraph(cell, text: str, size: int = 9, italic: bool = False) -> None:
    """Body text inside a panel — used for boilerplate clauses."""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.font.name = BODY_FONT
    run.font.color.rgb = BODY_DARK
    p.paragraph_format.space_after = Pt(4)


def section_field_row(
    cell,
    label: str,
    placeholder: str,
    *,
    label_width_cm: float = 4.0,
    field_width_cm: float = 13.0,
) -> None:
    """One row: red label on the left, white-bordered input on the
    right. Nested inside the section panel cell."""
    nested = cell.add_table(rows=1, cols=2)
    nested.autofit = False
    label_cell = nested.rows[0].cells[0]
    field_cell = nested.rows[0].cells[1]
    label_cell.width = Cm(label_width_cm)
    field_cell.width = Cm(field_width_cm)
    label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    field_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    lr = label_cell.paragraphs[0].add_run(label)
    lr.font.color.rgb = LABEL_RED
    lr.font.size = Pt(9)
    lr.font.name = BODY_FONT

    _set_cell_shading(field_cell, INPUT_BG)
    _set_cell_borders(field_cell)
    _set_cell_margin(field_cell, top=60, bottom=60, left=100, right=100)
    fr = field_cell.paragraphs[0].add_run(placeholder)
    fr.font.size = Pt(10)
    fr.font.name = BODY_FONT
    fr.font.color.rgb = BODY_DARK

    _strip_table_borders(nested)


def section_field_pair(
    cell,
    label_a: str,
    placeholder_a: str,
    label_b: str,
    placeholder_b: str,
) -> None:
    """Two label+input pairs side by side. Mirrors the Name/Position,
    Signature/Date, and Category/Cert-no rows on the Churches forms."""
    nested = cell.add_table(rows=1, cols=4)
    nested.autofit = False
    cells = nested.rows[0].cells
    cells[0].width = Cm(2.6)
    cells[1].width = Cm(6.4)
    cells[2].width = Cm(2.4)
    cells[3].width = Cm(6.1)
    for c in cells:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for label_idx, value_idx, label, placeholder in [
        (0, 1, label_a, placeholder_a),
        (2, 3, label_b, placeholder_b),
    ]:
        lr = cells[label_idx].paragraphs[0].add_run(label)
        lr.font.color.rgb = LABEL_RED
        lr.font.size = Pt(9)
        lr.font.name = BODY_FONT

        _set_cell_shading(cells[value_idx], INPUT_BG)
        _set_cell_borders(cells[value_idx])
        _set_cell_margin(cells[value_idx], top=60, bottom=60, left=100, right=100)
        fr = cells[value_idx].paragraphs[0].add_run(placeholder)
        fr.font.size = Pt(10)
        fr.font.name = BODY_FONT
        fr.font.color.rgb = BODY_DARK

    _strip_table_borders(nested)


def section_freeform(cell, placeholder: str, rows: int = 3) -> None:
    """Multi-line input box — white background, thin grey border, no
    label (the label sits above as a section title or prompt)."""
    nested = cell.add_table(rows=1, cols=1)
    nested.autofit = False
    inner = nested.rows[0].cells[0]
    inner.width = Cm(17.0)
    _set_cell_shading(inner, INPUT_BG)
    _set_cell_borders(inner)
    _set_cell_margin(inner, top=80, bottom=80, left=120, right=120)
    first = inner.paragraphs[0]
    first.add_run(placeholder).font.size = Pt(10)
    for _ in range(rows - 1):
        inner.add_paragraph()
    _strip_table_borders(nested)


def section_inline_check(cell, label: str) -> None:
    """Single ☐ + label paragraph for the System Examinations
    checklist (6 items on commissioning page 1)."""
    p = cell.add_paragraph()
    box = p.add_run(CHECKBOX_EMPTY + "  ")
    box.font.size = Pt(11)
    box.font.name = BODY_FONT
    lr = p.add_run(label)
    lr.font.size = Pt(9)
    lr.font.name = BODY_FONT
    lr.font.color.rgb = BODY_DARK
    p.paragraph_format.space_after = Pt(2)


# ── Page setup ───────────────────────────────────────────────────────

def set_a4_margins(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10)


# ── A056 — Installation Certificate ─────────────────────────────────

def build_installation_cert(out: Path) -> None:
    doc = Document()
    set_a4_margins(doc)
    set_default_font(doc)

    add_logo(doc)
    add_title(doc, "A056", "Fire Alarm Installation Certificate")

    panel = section_open(doc, "Certification of Installation for the Fire Alarm System at:")
    section_field_row(panel, "Address", "[Site Address]")
    section_field_row(panel, "Postcode", "[Site Postcode]", field_width_cm=6.0)

    section_body_paragraph(
        panel,
        "I/We being the competent person(s) responsible (as indicated by my/our "
        "signature(s) below) for the installation of the Fire Alarm System, particulars "
        "of which are set out below, CERTIFY that the said installation for which "
        "I/we have been responsible complies to the best of my/our knowledge and belief "
        "with the specification described below and with the recommendations of Section 4 "
        "of BS 5839-1 except for the variations, if any, stated in this certificate.",
    )

    section_field_pair(panel, "Name", "[Engineer Name]", "Position", "[Engineer Position]")
    section_field_pair(panel, "Signature", "[Engineer Signature]", "Date", "[Date]")
    section_field_row(
        panel,
        "For and behalf of",
        "[Company Name], [Company Address]",
    )
    section_field_pair(
        panel,
        "Category of system",
        "[Category of System]",
        "Certificate no",
        "[Certificate Number]",
    )

    add_spacer(doc, pt=2)
    panel2 = section_open(
        doc, "The extent of liability of the signatory is limited to the system described below:"
    )
    section_freeform(panel2, "[Extent of Liability]", rows=3)

    add_spacer(doc, pt=2)
    panel3 = section_open(
        doc,
        "Agreed variations from the specification and/or Section 4 of BS 5839-1 "
        "(see clause 7). Must be agreed by all parties",
    )
    section_freeform(panel3, "[Agreed Variations]", rows=3)

    add_spacer(doc, pt=2)
    panel4 = section_open(doc, "Inspection and testing of wiring systems")
    section_body_paragraph(
        panel4,
        "Unless supplied by others, the 'as fitted' drawings have been supplied to the "
        "person responsible for commissioning the system (see 36.2m of the current "
        "BS 5839-1). Test results have been recorded above.",
    )

    add_spacer(doc, pt=2)
    panel5 = section_open(doc, "Agreed variations from the specification — continuation page")
    section_freeform(panel5, "[Variations Continuation]", rows=10)

    add_spacer(doc, pt=4)
    add_submit_footer(doc)
    doc.save(out)


# ── A051 — Commissioning Certificate (3 pages) ──────────────────────

def build_commissioning_cert(out: Path) -> None:
    doc = Document()
    set_a4_margins(doc)
    set_default_font(doc)

    # ── Page 1 ──────────────────────────────────────────────────────
    add_logo(doc)
    add_title(doc, "A051", "Fire Alarm Commissioning Certificate")

    panel_client = section_open(doc, "Details of Client")
    section_field_row(panel_client, "Name", "[Customer Name]")
    section_field_row(panel_client, "Address", "[Customer Address]")
    section_field_row(panel_client, "Postcode", "[Customer Postcode]", field_width_cm=6.0)

    add_spacer(doc, pt=2)
    panel_sys = section_open(doc, "Details of the Fire Alarm and Detection System")
    section_field_row(panel_sys, "Address", "[Site Address]")
    section_field_row(panel_sys, "Extent of system covered by this", "[Extent of System]")
    section_field_pair(
        panel_sys,
        "The system is",
        "[New / Modification]",
        "Category",
        "[Category of System]",
    )

    add_spacer(doc, pt=2)
    panel_examine = section_open(doc, "System Examinations and Recommendations")
    section_body_paragraph(
        panel_examine,
        "Tick boxes or insert N/A (not applicable, as appropriate)",
        size=8,
        italic=True,
    )
    for label in [
        "All equipment operates correctly",
        "Installation work is, as far as can reasonably be ascertained, of an acceptable standard",
        "The entire system has been inspected and tested in accordance with the "
        "recommendations of clause 39.2c of the current BS 5839",
        "The system performs as required by the specification prepared by",
        "Taking into account the guidance in section 3 of the current BS 5839-1. I/we "
        "have not identified any obvious potential for any unacceptable rate of false alarms",
        "The documentation described in clause 40 of the standard has been provided to the user",
    ]:
        section_inline_check(panel_examine, label)

    section_field_row(
        panel_examine,
        "Specification prepared by",
        "[Specifier]",
        label_width_cm=5.0,
        field_width_cm=12.0,
    )
    section_field_pair(
        panel_examine,
        "Soak test period (weeks)",
        "[Soak Test Weeks]",
        "or",
        "[N/A]",
    )

    section_field_row(
        panel_examine,
        "Outstanding work before/after the system becomes operational",
        "[Outstanding Work]",
        label_width_cm=6.5,
        field_width_cm=10.5,
    )
    section_field_row(
        panel_examine,
        "Potential causes of false alarm at next service",
        "[False Alarm Risks]",
        label_width_cm=6.5,
        field_width_cm=10.5,
    )

    add_spacer(doc, pt=2)
    panel_cert = section_open(doc, "Certificate of Commissioning")
    section_body_paragraph(
        panel_cert,
        "I/We being the competent person(s) responsible (as indicated by my/our "
        "signature(s) below) for the commissioning of the fire alarm system, particulars "
        "of which are set out above, certify that the said work for which I/we have been "
        "responsible complies to the best of my/our knowledge and belief with the "
        "recommendations of clause 39 of BS 5839-1:2025, except for the variations, if any, "
        "stated in this certificate.",
    )
    section_field_row(
        panel_cert,
        "Variations from clause 39",
        "[Cl 39 Variations]",
        label_width_cm=5.0,
        field_width_cm=12.0,
    )
    section_field_pair(
        panel_cert,
        "Commissioning Engineer Name",
        "[Engineer Name]",
        "Position",
        "[Engineer Position]",
    )
    section_field_pair(panel_cert, "Signature", "[Engineer Signature]", "Date", "[Date]")

    add_spacer(doc, pt=2)
    panel_org = section_open(doc, "Particulars of the Organisation Commissioning the System")
    section_field_row(panel_org, "Organisation", "[Company Name], [Company Address]")
    section_field_pair(
        panel_org,
        "Design certificate no",
        "[Design Cert Number]",
        "Installation certificate no",
        "[Installation Cert Number]",
    )
    section_field_pair(
        panel_org,
        "Design drawings no",
        "[Design Drawings]",
        "As fitted drawings no",
        "[As Fitted Drawings]",
    )

    # ── Page 2 — 33-item checklist ──────────────────────────────────
    doc.add_page_break()
    add_logo(doc)
    add_title(doc, "A051", "Fire Alarm Commissioning Certificate")
    add_spacer(doc, pt=2)

    items = COMMISSIONING_CHECKLIST_ITEMS
    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.style = "Table Grid"
    widths = [Cm(0.9), Cm(13.2), Cm(1.0), Cm(1.0), Cm(1.2)]
    for col_idx, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[col_idx].width = w

    # Header row
    header = tbl.rows[0].cells
    for cell, label in zip(header, ["Item", "Description", "Y", "N", "N/A"]):
        _set_cell_shading(cell, CHECKLIST_HDR)
        _set_cell_margin(cell, top=80, bottom=80, left=80, right=80)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        run.font.name = BODY_FONT
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if label in ("Y", "N", "N/A", "Item")
            else WD_ALIGN_PARAGRAPH.LEFT
        )

    for row_idx, (item_num, description) in enumerate(items, start=1):
        cells = tbl.rows[row_idx].cells
        num_p = cells[0].paragraphs[0]
        num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = num_p.add_run(str(item_num))
        nr.font.size = Pt(9)
        nr.font.name = BODY_FONT

        desc = cells[1].paragraphs[0].add_run(description)
        desc.font.size = Pt(9)
        desc.font.name = BODY_FONT
        desc.font.color.rgb = BODY_DARK

        for box_col in (2, 3, 4):
            p = cells[box_col].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(CHECKBOX_EMPTY)
            r.font.size = Pt(11)
            r.font.name = BODY_FONT

    # ── Page 3 — incomplete work + sign-off ─────────────────────────
    doc.add_page_break()
    add_logo(doc)
    add_title(doc, "A051", "Fire Alarm Commissioning Certificate")

    panel_inc = section_open(
        doc,
        "The following work could not be completed for reasons beyond the control of the company:",
    )
    section_field_row(
        panel_inc, "Details", "[Incomplete Work Details]", label_width_cm=2.0, field_width_cm=15.0
    )
    section_field_row(
        panel_inc, "Reasons", "[Incomplete Work Reasons]", label_width_cm=2.0, field_width_cm=15.0
    )

    add_spacer(doc, pt=2)
    panel_visit = section_open(
        doc,
        "A further visit is required by an engineer to complete the installation "
        "(type N/A if not needed):",
    )
    section_freeform(panel_visit, "[Further Visit Required]", rows=4)

    add_spacer(doc, pt=2)
    panel_signoff = section_open(doc, None)
    section_body_paragraph(
        panel_signoff,
        "I certify that the above checks and functional tests have been completed in "
        "accordance with the equipment manufacturer's recommendations and/or applicable "
        "codes of practice.",
    )
    section_field_pair(
        panel_signoff,
        "Signature",
        "[Engineer Signature]",
        "Print",
        "[Engineer Name]",
    )

    add_spacer(doc, pt=4)
    add_submit_footer(doc)
    doc.save(out)


# ── A038 — Acceptance Certificate ───────────────────────────────────

def build_acceptance_cert(out: Path) -> None:
    doc = Document()
    set_a4_margins(doc)
    set_default_font(doc)

    add_logo(doc)
    add_title(doc, "A038", "Fire Alarm Acceptance Certificate")

    panel = section_open(
        doc, "Certificate of Installation (to BS 5839-1:2025) for the fire alarm system at:"
    )
    section_field_row(panel, "Address", "[Site Address]")
    section_field_row(panel, "Postcode", "[Site Postcode]", field_width_cm=6.0)

    section_body_paragraph(
        panel,
        "I/We being the competent person(s) responsible (as indicated by my/our "
        "signatures below) for the acceptance of the fire alarm system, particulars "
        "of which are set out below, ACCEPT the system on behalf of:",
    )
    section_field_pair(
        panel,
        "Customer Name",
        "[Customer Name]",
        "Position",
        "[Customer Position]",
    )
    section_field_pair(panel, "Signature", "[Customer Signature]", "Date", "[Date]")
    section_field_row(panel, "For and on behalf of", "[Customer Organisation]")

    add_spacer(doc, pt=2)
    panel2 = section_open(
        doc, "The extent of the liability of the signatory is limited to the system described below:"
    )
    section_field_row(
        panel2,
        "Extent of system covered by this certificate",
        "[Extent of System]",
        label_width_cm=6.5,
        field_width_cm=10.5,
    )

    add_spacer(doc, pt=2)
    panel3 = section_open(doc, "The following work is required before the system can be accepted:")
    section_freeform(panel3, "[Work Required]", rows=2)

    add_spacer(doc, pt=2)
    panel4 = section_open(doc, "Variations from the recommendations of clause 39 of BS 5839-1:2025")
    section_freeform(panel4, "[Cl 39 Variations]", rows=4)

    add_spacer(doc, pt=2)
    panel5 = section_open(
        doc, "Persons trained on the use of the fire alarm system, and how to prevent false alarms"
    )
    section_field_pair(panel5, "Name", "[Trained Person 1]", "Name", "[Trained Person 2]")
    section_field_pair(panel5, "Name", "[Trained Person 3]", "Name", "[Trained Person 4]")

    add_spacer(doc, pt=2)
    panel6 = section_open(doc, None)
    section_body_paragraph(panel6, "- All installation work appears to be satisfactory.")
    section_body_paragraph(panel6, "- The system is capable of giving a fire alarm signal.")
    section_body_paragraph(
        panel6,
        "- The facility for remote transmission of alarms to an alarm receiving centre "
        "operates correctly.",
    )

    add_spacer(doc, pt=2)
    panel7 = section_open(doc, "The following documents have been provided to the purchaser or user:")
    for line in [
        "- 'As fitted' drawings.",
        "- Operating and maintenance instructions.",
        "- Certificates of design, installation and commissioning.",
        "- A log book.",
        "- Sufficient representatives of the user have been properly instructed in the "
        "use of the system; including, at least, all means of triggering fire signals, "
        "silencing and resetting the system and avoidance of false alarms.",
        "- All relevant tests, defined in the purchasing specification, have been witnessed.",
    ]:
        section_body_paragraph(panel7, line)

    add_spacer(doc, pt=4)
    add_submit_footer(doc)
    doc.save(out)


# ── A058 — Battery Calculation ──────────────────────────────────────

def build_battery_calc(out: Path) -> None:
    doc = Document()
    set_a4_margins(doc)
    set_default_font(doc)

    add_logo(doc)
    add_title(doc, "A058", "Customer Fire Alarm Battery Calculation")

    panel_addr = section_open(doc, "Fire alarm battery calculations for")
    section_field_row(panel_addr, "Address", "[Site Address]")
    section_field_row(panel_addr, "Postcode", "[Site Postcode]", field_width_cm=6.0)

    add_spacer(doc, pt=2)
    panel_calc = section_open(doc, "Battery calculation")
    section_field_row(panel_calc, "Job number", "[Job Number]")
    section_field_pair(
        panel_calc,
        "Standby current 1 (A)",
        "[Standby Current]",
        "Standby time (hrs)",
        "[Standby Hours]",
    )
    section_field_pair(
        panel_calc,
        "+ Alarm current 2 (A)",
        "[Alarm Current]",
        "=  Sub-total (Ah)",
        "[Battery Subtotal]",
    )
    section_field_row(
        panel_calc,
        "× 1.25 = Minimum battery capacity (Ah)",
        "[Min Battery Capacity]",
        label_width_cm=7.0,
        field_width_cm=10.0,
    )
    section_body_paragraph(
        panel_calc,
        "Round up to the nearest standard battery size in Ah (e.g. 5.6 → 7 Ah).",
        size=8,
        italic=True,
    )
    section_field_pair(
        panel_calc,
        "Design battery size (Ah)",
        "[Design Battery]",
        "Installed battery size (Ah)",
        "[Installed Battery]",
    )

    # Warning callout — emphasised red text on the panel
    p = panel_calc.add_paragraph()
    wr = p.add_run(
        "If the calculation requires larger batteries than provided, please contact "
        "your PM or line manager."
    )
    wr.bold = True
    wr.font.size = Pt(9)
    wr.font.color.rgb = LABEL_RED
    wr.font.name = BODY_FONT

    add_spacer(doc, pt=2)
    panel_loops = section_open(doc, "Fire alarm loop calculations for")
    section_field_pair(
        panel_loops,
        "Panel located",
        "[Panel Location]",
        "Number of loops",
        "[Loop Count]",
    )
    section_body_paragraph(
        panel_loops,
        "For systems with multiple control panels, please complete an additional A058 "
        "sheet per panel.",
        size=8,
        italic=True,
    )

    add_spacer(doc, pt=2)
    panel_test = section_open(doc, "Test carried out by")
    section_field_pair(panel_test, "Name", "[Test Engineer Name]", "Test meter", "[Test Meter Model]")
    section_field_pair(
        panel_test, "Signature", "[Test Engineer Signature]", "Serial number", "[Test Meter Serial]"
    )
    section_field_row(panel_test, "Date", "[Test Date]", field_width_cm=6.0)

    add_spacer(doc, pt=4)
    add_submit_footer(doc)
    doc.save(out)


# 33-item commissioning checklist — verbatim from BS 5839-1 §39 as
# reflected in the A051 source form.
COMMISSIONING_CHECKLIST_ITEMS: list[tuple[int, str]] = [
    (1,  "The system complies with the original specification / design and the use of the building has not changed."),
    (2,  "The 'as fitted' drawing accurately reflects the building structure (any changes recorded and passed to PM)."),
    (3,  "System has been installed to meet the requirements of category L1 / L2 / L3 / L4 / L5 / P1 / P2 / M (state)."),
    (4,  "Variations to the defined category have been identified and the schedule of variations agreed by the client and a nominated designer."),
    (5,  "Cables meet requirements for standard / enhanced / mixed (BS 5839-1 §26.2)."),
    (6,  "Cables are segregated as required and suitably supported (where visibly checked, §26.2)."),
    (7,  "Cables are mechanically protected as required where necessary (§26.2)."),
    (8,  "Junction boxes correctly labelled and identified on drawings (where visibly checked)."),
    (9,  "All cable insulation and continuity resistance measurements are logged."),
    (10, "All cable penetrations are sleeved and fire-stopped (where visibly checked)."),
    (11, "Mains supply is dedicated, key-switched, correctly fused and labelled 'fire alarm — do not switch off' (red, accessible with a special tool only)."),
    (12, "Mains supply identified at ALL distribution boards with a 'fire alarm — do not switch off' label in red."),
    (13, "230V suppliers' installation covered by a certificate (request the EICR). If not seen, log variation as '230VAC supply test records not seen during commissioning BS 5839-1 §38.2c.'"),
    (14, "Standby battery verification calculation has been carried out for ALL panels / power supplies. If calculation requires larger batteries than provided, contact PM or line manager."),
    (15, "All batteries are clearly marked and labelled with date of installation (§25.4)."),
    (16, "Field wiring is labelled and correctly terminated in all control and ancillary equipment."),
    (17, "Detector removal fault indication has been checked and tested (§12.2.1)."),
    (18, "Short circuit fault indication has been checked and tested."),
    (19, "ALL detection, MCPs, warning and ancillary devices have been tested for control operation and results recorded."),
    (20, "Cause and effect on the system has been tested and results recorded."),
    (21, "Provision of sounder circuits is appropriate (§16.2)."),
    (22, "Sound pressure levels have been checked and recorded (§16.2)."),
    (23, "Detector type and spacing is appropriate to the system category (§22.2)."),
    (24, "MCPs are located correctly and travel distances do not exceed 45 m (§22.2)."),
    (25, "Remote signalling has been checked and tested for correct operation to ARC (fire and fault)."),
    (26, "Radio signal strength (where applicable) exceeds manufacturer's minimum requirements."),
    (27, "Cause and effects have been checked and verified."),
    (28, "Zone charts have been fitted in appropriate locations and with the correct orientation (e.g. adjacent to control equipment and repeaters); search distances do not exceed limits; emergency lighting above."),
    (29, "'As fitted' drawings are complete and updated where required including cable type and sizes, cable routing, 230V supply, location of all MCPs / detectors / sounders / isolators."),
    (30, "User handbook and operating instructions have been issued to the 'responsible person'."),
    (31, "The 'responsible person' has been adequately trained in the use of the fire alarm system; details recorded."),
    (32, "Premises have been left in a tidy condition and all surplus materials and equipment removed from site."),
    (33, "Signage — correct signage has been installed (call points and fire action signs)."),
]


def main() -> int:
    if not ASSETS.exists():
        print(f"assets/ not found: {ASSETS}", file=sys.stderr)
        return 1
    targets = [
        ("bs5839-installation-cert-template.docx",  build_installation_cert),
        ("bs5839-commissioning-cert-template.docx", build_commissioning_cert),
        ("bs5839-acceptance-cert-template.docx",    build_acceptance_cert),
        ("bs5839-battery-calc-template.docx",       build_battery_calc),
    ]
    for filename, builder in targets:
        out = ASSETS / filename
        builder(out)
        print(f"  ✓ {filename:<48} {out.stat().st_size:>6} bytes")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
